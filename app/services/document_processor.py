# app/services/document_processor.py

import openai
# from app.core.config import settings
from typing import Dict, Any, Optional, Union
from io import BytesIO
from docx import Document
import pytesseract
from PIL import Image, UnidentifiedImageError
import pyheif
import pdfplumber
import re
import json
import logging
import os

# Initialize OpenAI API key
# openai.api_key = settings.OPENAI_API_KEY

# Configure logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

def extract_text_from_filepath(file_path: str) -> Optional[str]:
    """
    Extracts text from an image (PNG, JPG, HEIC) or PDF/DOCX file using OCR and document parsing.

    Args:
        file_path (str): The path to the file.

    Returns:
        Optional[str]: The extracted text, or None if extraction fails.
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    text = ""

    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return None

    try:
        if file_extension in {".png", ".jpg", ".jpeg"}:
            # Handle image files
            logger.info(f"Extracting text from image: {file_path}")
            try:
                image = Image.open(file_path)
            except UnidentifiedImageError as e:
                logger.error(f"Unable to open image file: {file_path}. Error: {e}")
                return None

            text = pytesseract.image_to_string(image)
            if not text.strip():
                logger.warning(f"No text extracted from image: {file_path}")
                return None

        elif file_extension == ".heic":
            # Handle HEIC (iPhone photo format)
            logger.info(f"Extracting text from HEIC image: {file_path}")
            try:
                heif_file = pyheif.read(file_path)
                image = Image.frombytes(
                    heif_file.mode,
                    heif_file.size,
                    heif_file.data,
                    "raw",
                    heif_file.mode,
                    heif_file.stride,
                )
            except Exception as e:
                logger.error(f"Unable to process HEIC file: {file_path}. Error: {e}")
                return None

            text = pytesseract.image_to_string(image)
            if not text.strip():
                logger.warning(f"No text extracted from HEIC image: {file_path}")
                return None

        elif file_extension == ".pdf":
            # Handle PDF files
            logger.info(f"Extracting text from PDF: {file_path}")
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_number, page in enumerate(pdf.pages, start=1):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            logger.warning(f"No text extracted from page {page_number} of PDF: {file_path}")
            except Exception as e:
                logger.error(f"Error processing PDF file: {file_path}. Error: {e}")
                return None

            if not text.strip():
                logger.warning(f"No text extracted from PDF: {file_path}")
                return None

        elif file_extension == ".docx":
            # Handle DOCX files
            logger.info(f"Extracting text from DOCX file: {file_path}")
            try:
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
            except Exception as e:
                logger.error(f"Error processing DOCX file: {file_path}. Error: {e}")
                return None

            if not text.strip():
                logger.warning(f"No text extracted from DOCX file: {file_path}")
                return None

        else:
            # Unsupported file type
            logger.warning(f"Unsupported file type: {file_extension}")
            return None

    except Exception as e:
        logger.error(f"Unexpected error during text extraction: {e}")
        return None

    return text.strip()

def extract_text_from_file(file: Union[BytesIO, 'File'], filename: str) -> Optional[str]:
    """
    Extracts text from an image (PNG, JPG, HEIC), PDF, or DOCX file using OCR and document parsing.

    Args:
        file (BytesIO or File): The file-like object containing the file data.
        filename (str): The name of the file, used to determine the file type.

    Returns:
        Optional[str]: The extracted text, or None if extraction fails.
    """
    import os  # Ensure os is imported

    file_extension = os.path.splitext(filename)[1].lower()
    text = ""

    try:
        if file_extension in {".png", ".jpg", ".jpeg"}:
            # Handle image files
            logger.info(f"Extracting text from image: {filename}")
            try:
                image = Image.open(file)
            except UnidentifiedImageError as e:
                logger.error(f"Unable to open image file: {filename}. Error: {e}")
                return None

            text = pytesseract.image_to_string(image)

        elif file_extension == ".heic":
            # Handle HEIC (iPhone photo format)
            logger.info(f"Extracting text from HEIC image: {filename}")
            try:
                heif_file = pyheif.read(file.read())
                image = Image.frombytes(
                    heif_file.mode,
                    heif_file.size,
                    heif_file.data,
                    "raw",
                    heif_file.mode,
                    heif_file.stride,
                )
            except Exception as e:
                logger.error(f"Unable to process HEIC file: {filename}. Error: {e}")
                return None

            text = pytesseract.image_to_string(image)

        elif file_extension == ".pdf":
            # Handle PDF files
            logger.info(f"Extracting text from PDF: {filename}")
            try:
                with pdfplumber.open(file) as pdf:
                    for page_number, page in enumerate(pdf.pages, start=1):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            logger.warning(f"No text extracted from page {page_number} of PDF: {filename}")
            except Exception as e:
                logger.error(f"Error processing PDF file: {filename}. Error: {e}")
                return None

        elif file_extension == ".docx":
            # Handle DOCX files
            logger.info(f"Extracting text from DOCX file: {filename}")
            try:
                doc = Document(file)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
            except Exception as e:
                logger.error(f"Error processing DOCX file: {filename}. Error: {e}")
                return None

        else:
            # Unsupported file type
            logger.warning(f"Unsupported file type: {file_extension}")
            return None

        # Check if any text was extracted
        if not text.strip():
            logger.warning(f"No text extracted from file: {filename}")
            return None

    except Exception as e:
        logger.error(f"Unexpected error during text extraction: {e}")
        return None

    return text.strip()

def process_document(file_path: str, document_type: str) -> Dict[str, Any]:
    """
    Processes the document using OpenAI API to extract relevant data.

    Args:
        file_path (str): The path to the file.
        document_type (str): The type of the document (invoice, lease, receipt).

    Returns:
        Dict[str, Any]: The extracted data mapped to database schema fields.
    """
    text_content = extract_text_from_file(file_path)
    if not text_content.strip():
        logger.error("No text extracted from the document.")
        raise ValueError("No text extracted from the document.")

    prompt = generate_prompt(document_type, text_content)

    try:
        logger.info("Sending request to OpenAI API.")
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=500,
            temperature=0.2,
            n=1,
            stop=None,
        )
        logger.info("Received response from OpenAI API.")
        extracted_data = parse_openai_response(response.choices[0].text)
        return extracted_data
    except Exception as e:
        logger.error(f"Error during OpenAI processing: {str(e)}")
        raise

def generate_prompt(document_type: str, text_content: str) -> str:
    """
    Generates a prompt for OpenAI based on the document type.

    Args:
        document_type (str): The type of the document.
        text_content (str): The extracted text from the document.

    Returns:
        str: The prompt for OpenAI API.
    """
    prompts = {
        "invoice": (
            "Extract the following information from the invoice and provide the output in JSON format:\n"
            "- Invoice Number\n"
            "- Vendor Name\n"
            "- Amount\n"
            "- Invoice Date\n"
        ),
        "lease": (
            "Extract the following information from the lease agreement and provide the output in JSON format:\n"
            "- Tenant Name\n"
            "- Start Date\n"
            "- End Date\n"
            "- Monthly Rent\n"
        ),
        "receipt": (
            "Extract the following information from the receipt and provide the output in JSON format:\n"
            "- Merchant Name\n"
            "- Amount\n"
            "- Receipt Date\n"
        ),
    }

    base_prompt = prompts.get(document_type.lower())
    if not base_prompt:
        logger.warning(f"Unknown document type: {document_type}. Using generic extraction.")
        base_prompt = (
            "Extract key information from the following text and provide the output in JSON format.\n"
        )

    full_prompt = base_prompt + "\n" + text_content
    return full_prompt

def parse_openai_response(response_text: str) -> Dict[str, Any]:
    """
    Parses the OpenAI API response text into a structured dictionary.

    Args:
        response_text (str): The raw response text from OpenAI.

    Returns:
        Dict[str, Any]: The structured data.
    """
    try:
        # Attempt to extract JSON from the response
        json_str = re.search(r'\{.*\}', response_text, re.DOTALL).group()
        extracted_data = json.loads(json_str)
        logger.info("Successfully parsed OpenAI response.")
        return extracted_data
    except Exception as e:
        logger.error(f"Failed to parse OpenAI response: {str(e)}")
        raise ValueError("Failed to parse OpenAI response as JSON.")



def parse_custom_json(json_data: dict) -> dict:
    """
    Parses a JSON object with dynamic keys and stores its information in a structured format.
    
    Args:
        json_data (dict): The JSON data to parse.
    
    Returns:
        dict: A structured representation of the parsed data.
    """
    parsed_data = {}

    def recurse_json(data, parent_key=""):
        if isinstance(data, dict):
            for key, value in data.items():
                full_key = f"{parent_key}.{key}" if parent_key else key
                recurse_json(value, full_key)
        elif isinstance(data, list):
            for index, item in enumerate(data):
                full_key = f"{parent_key}[{index}]"
                recurse_json(item, full_key)
        else:
            parsed_data[parent_key] = value

    recurse_json(json_data)
    return parsed_data

if __name__ == "__main__":
    # text = extract_text_from_file("/Users/tommyqu/Desktop/Resume/Thomas Qu - Resume.pdf")
    text = extract_text_from_file("/Users/tommyqu/Downloads/receipt1.jpeg")
    print(text)
    print("="*70)
    text = extract_text_from_file("/Users/tommyqu/Downloads/receipt2.jpeg")
    print("HELLo")
    print(text)
    print("HELLo")